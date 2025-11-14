#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档解析脚本
使用python-docx库提取.docx文件的文本内容
"""

import sys
import os
from docx import Document


def parse_docx(file_path):
    """
    解析Word文档，提取所有文本内容

    Args:
        file_path: docx文件路径

    Returns:
        str: 提取的文本内容
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 加载文档
        doc = Document(file_path)

        # 提取所有段落文本
        full_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # 跳过空段落
                full_text.append(text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(' | '.join(row_text))

        # 合并所有文本
        result = '\n'.join(full_text)

        return result

    except Exception as e:
        # 返回错误信息
        return f"ERROR: {str(e)}"


def main():
    """
    命令行入口
    接收文件路径作为参数，输出解析结果
    """
    if len(sys.argv) < 2:
        print("ERROR: 请提供文件路径作为参数", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]

    # 解析文档
    result = parse_docx(file_path)

    # 输出结果（通过stdout）
    print(result)

    # 如果是错误，返回非0退出码
    if result.startswith("ERROR:"):
        sys.exit(1)


if __name__ == "__main__":
    main()
